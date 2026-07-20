"""Generate the Upande Webstore architecture description as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# Title page
title = doc.add_heading("Upande Webstore", level=0)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run("E-commerce Webstore & Customer Portal for ERPNext v16\nArchitecture Description")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
meta = doc.add_paragraph()
r = meta.add_run("Document 1 of the design series  |  20 July 2026  |  Status: Approved in design review")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

heading("1. Overview", 1)
para(
    "The Upande Webstore is a combined e-commerce storefront and customer portal built as a "
    "single custom Frappe application installed on an ERPNext v16 bench. It serves both retail "
    "(B2C) and business (B2B) customers: visitors browse a public catalog with standard selling "
    "prices, while logged-in business customers see their own ERPNext price lists and pricing "
    "rules. There is no online payment at launch; checkout produces an ERPNext Quotation which "
    "the sales team reviews and converts to a Sales Order, with payment handled offline "
    "(invoice, bank transfer, or credit terms)."
)
para(
    "The customer portal covers the full commercial relationship: quotations (view and accept), "
    "order history and status, invoices (view and download), outstanding balance and account "
    "statement, and support tickets backed by ERPNext Issues."
)

heading("2. Key Decisions", 1)
decisions = [
    ("Platform: ", "Custom Frappe app on the ERPNext v16 bench (not the frappe/webshop app, not a "
     "headless frontend). This gives the tightest integration with ERPNext pricing, customers, and "
     "documents, a single deployment, and full v16 compatibility."),
    ("Frontend: ", "Server-rendered Jinja portal pages with lightweight JavaScript modules for the "
     "cart, variant selection, and quantity updates. SEO-friendly product pages out of the box, no "
     "separate build pipeline beyond Frappe's own asset bundling."),
    ("Audience: ", "Both B2C and B2B. Guests browse freely; an account is required to place an "
     "order. Signup automatically creates an ERPNext Customer and portal user."),
    ("Order flow: ", "Quotation-first. Every web order is created as a Quotation visible in the "
     "customer's portal; the sales team converts it to a Sales Order."),
    ("Payments: ", "None at launch. Payment is handled offline."),
    ("Catalog: ", "Customer-specific B2B pricing, item variants, stock availability display, "
     "product search with category filters, and a per-user wishlist."),
    ("Out-of-stock items: ", "Not orderable. Add-to-cart is disabled in the UI and rejected by "
     "the API; checkout re-validates stock for every line before creating the quotation."),
]
for lead, rest in decisions:
    bullet(rest, bold_lead=lead)

heading("3. Application Structure", 1)
para("One custom Frappe app, upande_webstore, containing:")
bullet("routes for the storefront and portal pages, rendered with Jinja templates and "
       "Python page controllers.", bold_lead="www/ ")
bullet("module of whitelisted REST endpoints: cart operations, variant resolution, and "
       "stock/price lookups.", bold_lead="api/ ")
bullet("Public JavaScript and CSS bundled through Frappe's build system "
       "(upande_webstore.bundle.js).", bold_lead="Assets: ")

heading("4. Data Model", 1)
para(
    "New DocTypes are kept to a minimum; ERPNext remains the system of record for everything "
    "transactional."
)

heading("4.1 New DocTypes", 2)
numbered("(Single) — default price list for guests/B2C, default warehouse(s) for stock "
         "display, default Customer Group and Territory for signups, quotation validity days, and "
         "a toggle for showing exact quantity versus in/out of stock.",
         bold_lead="Webstore Settings ")
numbered("— the published-item record. Links to an ERPNext Item (template or standalone) "
         "and carries web-specific fields: slug/route, web title, short and long description, "
         "images, published flag, featured flag, and category (Item Group link) for filtering. "
         "This mirrors the role of the webshop app's Website Item, but is owned by this app and "
         "native to v16.",
         bold_lead="Webstore Product ")
numbered("— server-side cart: owner (user), status (Open / Ordered / Abandoned), and a "
         "child table of items (Item, quantity, resolved rate). One open cart per user.",
         bold_lead="Webstore Cart ")
numbered("— one per user; a child table of saved Webstore Products with the date added. "
         "Backs the wishlist heart toggle on product cards/pages and the /wishlist page.",
         bold_lead="Webstore Wishlist ")

heading("4.2 Stock ERPNext DocTypes Used As-Is", 2)
para(
    "Customer, Contact, Item and Item variants, Price List and Item Price, Pricing Rule, Bin "
    "(stock levels), Quotation, Sales Order, Sales Invoice, Payment Entry, and Issue. Checkout "
    "maps the cart to a Quotation linked to the customer's Contact, so it appears in their "
    "portal immediately."
)

heading("5. Pricing Resolution", 1)
para(
    "A single shared server-side function resolves prices for product pages, the cart, and "
    "checkout:"
)
numbered("If the logged-in user's Customer has a default price list, use it, then apply Pricing "
         "Rules via ERPNext's own pricing machinery (get_price_list_rate and related utilities).")
numbered("Otherwise, use the guest price list configured in Webstore Settings.")
para(
    "Prices are never trusted from the client. Every add-to-cart, cart update, and checkout "
    "re-resolves prices server-side.", bold_lead="Security invariant: "
)

heading("6. Order Flow", 1)
numbered("Visitor browses the catalog (search, category filters, variant selection, stock "
         "availability).")
numbered("To order, the visitor signs in or signs up; signup creates an ERPNext Customer, "
         "Contact, and portal user automatically.")
numbered("The customer builds a server-side cart and checks out.")
numbered("Checkout creates a Quotation from the cart with server-resolved prices.")
numbered("The sales team reviews the Quotation and converts it to a Sales Order; the customer "
         "can view and accept quotations from the portal.")
numbered("Invoicing, payment (offline), and fulfilment proceed in ERPNext as normal, all "
         "visible to the customer in the portal.")

doc.save("/home/austin/vscodeProjects/upande_webstore/docs/01-architecture-description.docx")
print("saved")
