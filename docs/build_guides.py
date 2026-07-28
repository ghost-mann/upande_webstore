"""Generate the Upande Webstore user, developer and customisation guides as .docx.

python-docx lives in the bench virtualenv, not system python, so run:
    ../../env/bin/python docs/build_guides.py
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))


def new_doc(title, subtitle):
	doc = Document()
	style = doc.styles["Normal"]
	style.font.name = "Calibri"
	style.font.size = Pt(11)
	doc.add_heading(title, level=0)
	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	r = p.add_run(subtitle)
	r.font.size = Pt(13)
	r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
	meta = doc.add_paragraph()
	r = meta.add_run("Upande Webstore for ERPNext v16  |  20 July 2026")
	r.font.size = Pt(9)
	r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
	return doc


def h(doc, text, level=1):
	doc.add_heading(text, level=level)


def p(doc, text, bold_lead=None):
	par = doc.add_paragraph()
	if bold_lead:
		r = par.add_run(bold_lead)
		r.bold = True
	par.add_run(text)
	return par


def bullet(doc, text, bold_lead=None):
	par = doc.add_paragraph(style="List Bullet")
	if bold_lead:
		r = par.add_run(bold_lead)
		r.bold = True
	par.add_run(text)
	return par


def num(doc, text, bold_lead=None):
	par = doc.add_paragraph(style="List Number")
	if bold_lead:
		r = par.add_run(bold_lead)
		r.bold = True
	par.add_run(text)
	return par


def code(doc, text):
	par = doc.add_paragraph()
	r = par.add_run(text)
	r.font.name = "Consolas"
	r.font.size = Pt(9)
	return par


# ======================================================================
# USER GUIDE
# ======================================================================
doc = new_doc("Upande Webstore — User Guide", "For customers and for the Upande sales & operations team")

h(doc, "Part 1 — For Customers", 1)

h(doc, "1. Creating your account", 2)
num(doc, "Open the store and click Sign up (or the Member login button if you already have credentials).")
num(doc, "Fill in your name, email and phone. If you are buying for a company, enter the company name — your account and all documents will be registered to it.")
num(doc, "You will receive an email to set your password. If it does not arrive, contact the sales team — they can set portal access for you directly.")
p(doc, "If your company already trades with Upande, ask the sales team to link a portal login to your existing customer account — your agreed price list, credit terms and full document history apply automatically.", bold_lead="Existing customers: ")

h(doc, "2. Browsing and searching the catalog", 2)
bullet(doc, "The Store page lists all published products with live prices and stock availability. Filter by category (Flowers, Coffee, Fresh Produce) in the sidebar, or use the search box.")
bullet(doc, "Press Ctrl+K (⌘K on Mac) or the search field in the top bar from any page to open the quick search — type at least two letters, use the arrow keys, and press Enter to open a product.")
bullet(doc, "A green dot means the item is in stock; a grey dot means it is currently unavailable and cannot be ordered. Products with options (e.g. sizes) show the price after you select a combination.")
bullet(doc, "If you are logged in and your company has an agreed price list, you will see your prices, marked “Your price”.")

h(doc, "3. Basket and requesting a quotation", 2)
num(doc, "On a product page choose a quantity and click Add to cart. The Basket in the top bar opens a side panel where you can adjust quantities or remove lines at any time.")
num(doc, "When ready, open the Basket and click “Review & request quotation”. On the checkout page choose a delivery address, optionally add your PO reference and notes, and click Request Quotation.")
num(doc, "Nothing is charged and no stock is committed at this point: your request creates a quotation which the sales team reviews and confirms — usually within one working day.")
p(doc, "Prices shown in the basket are always re-confirmed by the server when the quotation is created; quantities are checked against live stock at that moment.", bold_lead="Good to know: ")

h(doc, "4. Your portal", 2)
p(doc, "Click My Portal (or Member login) to reach your dashboard. It shows your outstanding balance, open quotations, and recent orders. The tabs along the top take you to:")
bullet(doc, "list of your quotations with status. Open one to review line items and totals, then Accept or Decline it. Accepting notifies the sales team, who convert it into a confirmed order.", bold_lead="Quotations — ")
bullet(doc, "confirmed sales orders with delivery and billing progress.", bold_lead="Orders — ")
bullet(doc, "all invoices with due dates and status (paid / unpaid / overdue). Each invoice can be downloaded as a PDF.", bold_lead="Invoices — ")
bullet(doc, "a date-filtered account statement with running balance, printable to PDF.", bold_lead="Statement — ")
bullet(doc, "raise general support tickets and follow the team's replies.", bold_lead="Support — ")
bullet(doc, "report a problem with a delivery: choose the type (damaged goods, short delivery, quality below grade, billing error, other), reference the order or invoice concerned, describe what happened and attach photos. Claims are tracked separately from support tickets.", bold_lead="Claims — ")
bullet(doc, "update your name and phone, manage delivery addresses, change your password, and log out.", bold_lead="Account — ")

h(doc, "Part 2 — For the Upande Sales & Operations Team", 1)

h(doc, "5. One-time configuration", 2)
p(doc, "In the desk, open Webstore Settings and set:")
bullet(doc, "Company, Guest Price List (what visitors see), and the warehouses whose stock counts toward availability.")
bullet(doc, "Default Customer Group and Territory applied to self-service signups.")
bullet(doc, "Quotation validity in days, stock display mode (badge or exact quantity), and the notification email(s) that receive new webstore quotations, acceptances, declines and claims.")

h(doc, "6. Publishing products", 2)
num(doc, "Create the Item in ERPNext as usual (templates with variants are supported), with an Item Price on the guest price list and any customer price lists.")
num(doc, "Create a Webstore Product: link the Item, write the web title and descriptions, attach a photo, set the category (Item Group), and tick Published.")
num(doc, "Tick Featured to place the product in the “Available this week” card on the store homepage.")
p(doc, "Products with zero stock across the configured warehouses remain visible but cannot be ordered. Customer-specific pricing needs no extra setup — the store reads the customer's default price list and ERPNext pricing rules.", bold_lead="Notes: ")

h(doc, "7. Handling webstore documents", 2)
bullet(doc, "arrive as submitted Quotations (order type “Shopping Cart”) with the customer's PO reference and notes in the Webstore section. You are notified by email. Review, adjust if needed, and use Create > Sales Order once agreed.", bold_lead="New quotation requests ")
bullet(doc, "set the quotation's Portal Status field and add a comment; you are notified. Convert accepted quotations promptly.", bold_lead="Customer accept/decline ")
bullet(doc, "appear as Issues with type “Claim”; ordinary support tickets have no type. Replies you send on the Issue are visible to the customer in their portal.", bold_lead="Claims ")
bullet(doc, "to give an existing ERPNext customer portal access, open one of their Contacts, set its User field to a website user with the Customer role.", bold_lead="Portal access ")

doc.save(os.path.join(HERE, "02-user-guide.docx"))
print("saved 02-user-guide.docx")

# ======================================================================
# DEVELOPER GUIDE
# ======================================================================
doc = new_doc("Upande Webstore — Developer Guide", "Architecture, stack, environment and maintenance handbook")

h(doc, "1. What this app is", 1)
p(doc, "upande_webstore is a single custom Frappe app for ERPNext v16 providing a B2B/B2C storefront and customer portal. Checkout is quotation-first (no online payments): the cart becomes a submitted ERPNext Quotation which the sales team converts to a Sales Order. ERPNext remains the system of record; the app adds four DocTypes and a thin service/API layer. See docs/01-architecture-description.docx and docs/superpowers/specs/ for the original design.")

h(doc, "2. Repository layout", 1)
code(doc, "upande_webstore/\n"
	"  api/        whitelisted endpoints: cart, checkout, account (signup/profile/address),\n"
	"              variants, wishlist, portal (accept/decline/invoice PDF), support (+claims), search\n"
	"  services/   business logic: settings, pricing, stock, catalog, portal (scoping),\n"
	"              portal_data, statement\n"
	"  upande_webstore/doctype/   Webstore Settings, Webstore Product, Webstore Cart(+Item),\n"
	"              Webstore Wishlist(+Item), Webstore Warehouse\n"
	"  www/        store, cart, signup, wishlist + www/portal/* pages (Jinja + controllers)\n"
	"  templates/  webstore_base.html (navbar/footer/dialogs), includes (portal hero/nav/macros),\n"
	"              generators/webstore_product.html\n"
	"  public/     js/webstore.bundle.ts, scss/webstore.bundle.scss, tailwind/input.css,\n"
	"              css/tailwind.css (built), fonts/, images/\n"
	"  tests/      integration tests + shared fixtures in tests/utils.py\n"
	"  setup/install.py   Quotation custom fields (after_install/after_migrate)")

h(doc, "3. Development environment", 1)
bullet(doc, "bench at ~/frappe-v16-bench (frappe & erpnext v16.27.0); dev site webstore.localhost (Administrator/admin).", bold_lead="Bench: ")
bullet(doc, "the bench has serve_default_site=kaitet.local on port 8002, so the webstore is served by a dedicated process: bench --site webstore.localhost serve --port 8003. Restart it after reboots; do not change the bench default.", bold_lead="Port 8003: ")
bullet(doc, "canonical working copy is ~/frappe-v16-bench/apps/upande_webstore; git origin is /home/austin/vscodeProjects/upande_webstore (receive.denyCurrentBranch=updateInstead, so pushing updates that checkout).", bold_lead="Git: ")
bullet(doc, "if working from a Flatpak-sandboxed editor, bench/DB commands must run on the host: flatpak-spawn --host bash -lc '<command>'. Also note bash -lc does not source nvm — source ~/.nvm/nvm.sh before bench build.", bold_lead="Sandbox: ")
bullet(doc, "demo portal login demo@upande.com / upande123; MariaDB root password on this dev machine is 'root'.", bold_lead="Dev credentials: ")

h(doc, "4. Frontend stack and build", 1)
bullet(doc, "shadcn/ui visual spec: white background, zinc neutrals (border #e4e4e7, muted #71717a), primary green #166534, Inter, radius 0.625rem, shadow-xs/sm. No emoji, no gradients.", bold_lead="Design system: ")
bullet(doc, "tokens and shadcn variables live in public/tailwind/input.css (imports tailwind theme+utilities and basecoat-css). Rebuild after template/theme changes:", bold_lead="Tailwind v4: ")
code(doc, "cd apps/upande_webstore   # npm install (first time)\nnpx @tailwindcss/cli -i upande_webstore/public/tailwind/input.css \\\n    -o upande_webstore/public/css/tailwind.css --minify   # add --watch in dev")
bullet(doc, "bespoke components (ws-* classes: cards, tables, hero, portal tabs, drawer, palette, toasts) live in public/scss/webstore.bundle.scss, compiled by bench build.", bold_lead="Component layer: ")
bullet(doc, "public/js/webstore.bundle.ts (frappe esbuild compiles .bundle.ts natively). Exposes window.webstore = { addToCart, toggleWishlist, refreshCartBadge, openCart, openPalette, call, toast }. Interactive components: ⌘K command palette (api.search.search_products), slide-over basket drawer, sonner-style toasts. Native <dialog> elements; Esc/backdrop close; prefers-reduced-motion respected.", bold_lead="TypeScript runtime: ")
code(doc, "bench build --app upande_webstore && bench --site webstore.localhost clear-website-cache")
bullet(doc, "hooks.py: web_include_css = [webstore.bundle.css, /assets/upande_webstore/css/tailwind.css], web_include_js = webstore.bundle.js. These load on ALL website pages of the site.", bold_lead="Includes: ")

h(doc, "5. Backend invariants", 1)
bullet(doc, "prices are NEVER trusted from the client. services/pricing.get_item_price() resolves via erpnext get_item_details (price list + pricing rules); the cart re-prices on every read/mutation; checkout re-resolves again.", bold_lead="Pricing: ")
bullet(doc, "out-of-stock items are not orderable — enforced in api/cart.add_item/update_qty AND re-validated in api/checkout.place_order (services/stock sums Bin.actual_qty over configured warehouses).", bold_lead="Stock: ")
bullet(doc, "every portal read goes through services/portal.py: get_current_customer() (session user → Contact → Customer), get_customer_docs() (always injects the customer filter) and assert_customer_doc() (403 on other customers' documents). Never query transactional doctypes in a portal page without these.", bold_lead="Isolation: ")
bullet(doc, "checkout and invoice-PDF rendering run under temporarily elevated context (frappe.set_user Administrator in try/finally) AFTER ownership checks, because ERPNext's account/print permission checks have no website-user path.", bold_lead="Elevation pattern: ")
bullet(doc, "claims are Issues with issue_type='Claim' (created idempotently); support lists exclude them in Python because SQL != misses NULL issue_type.", bold_lead="Claims: ")

h(doc, "6. ERPNext v16 gotchas (hard-won)", 1)
bullet(doc, "frappe.get_all rejects SQL strings like sum(x) or count(x) in fields — aggregate in Python or use query builder.")
bullet(doc, "WebsiteGenerator.autoname sets the scrubbed route name; combined with autoname='field:web_title' the field itself gets overwritten by _sync_autoname_field — do not add field: autoname to website generator doctypes.")
bullet(doc, "Override validate() with super().validate() in WebsiteGenerator subclasses or routes never get set.")
bullet(doc, "GET /?cmd=web_logout returns 403 (POST-only since v16) — log out via POST /api/method/logout with CSRF token.")
bullet(doc, "erpnext get_balance_on requires desk permissions — the portal computes balances/statements directly from GL Entry.")
bullet(doc, "@rate_limit decorates sign_up (20/hour/IP); frappe.rate_limiter.apply_for does not exist.")
bullet(doc, "Fresh sites: if erpnext's after_install did not complete, Customer role may keep desk_access=1 (breaking website-user creation) and Contact custom fields may be missing — re-run erpnext.setup.install.after_install.")
bullet(doc, "Stock Reconciliation on an empty site trips opening-entry account validation — test fixtures use Material Receipt/Issue Stock Entries instead.")
bullet(doc, "wkhtmltopdf needs a resolvable host for assets — site_config host_name is set to http://127.0.0.1:8003.")

h(doc, "7. Testing", 1)
code(doc, "bench --site webstore.localhost set-config allow_tests true   # once\nbench --site webstore.localhost run-tests --app upande_webstore  # full suite (220 tests)\nbench --site webstore.localhost run-tests --module upande_webstore.tests.test_cart")
p(doc, "The invoice-PDF test needs a running web server (bench serve) so wkhtmltopdf can fetch print assets; it fails with a connection error otherwise. Pure-function suites (test_theme colour math) need no site and run under plain unittest.", bold_lead="Two environment notes: ")
p(doc, "Shared fixtures live in upande_webstore/tests/utils.py (setup_webstore_settings, make_test_product, make_portal_user, make_item_price, set_stock, make_variant_template). Every portal feature has cross-customer isolation tests — keep that bar when adding endpoints.")

h(doc, "8. Common tasks", 1)
bullet(doc, "add the page under www/portal/ (controller calls portal_guard(route) first), set portal_active/portal_title in a hero block including webstore_portal_hero.html, add the tab in templates/includes/webstore_portal_nav.html, write an isolation test.", bold_lead="New portal page: ")
bullet(doc, "add a function in api/, decorate @frappe.whitelist (methods=['POST'] for mutations), call _require_login()/get_current_customer() as appropriate, never trust client prices/quantities beyond validation.", bold_lead="New API: ")
bullet(doc, "per-client changes are configuration, not code — use Webstore Settings (see the Customisation Manual); no rebuild or deploy needed. Only change the SCSS :root block in public/scss/webstore.bundle.scss when you are altering the shipped DEFAULT for every site, and remember blank settings must keep rendering identically. Derivation lives in upande_webstore/theme/color.py.", bold_lead="Theme change: ")
bullet(doc, "payments (M-Pesa/cards) were deliberately deferred: the seam is after invoice creation — add a payment gateway integration against Sales Invoice outstanding amounts; nothing in checkout needs to change.", bold_lead="Adding payments later: ")

doc.save(os.path.join(HERE, "03-developer-guide.docx"))
print("saved 03-developer-guide.docx")


# ======================================================================
# CUSTOMISATION MANUAL
# ======================================================================
doc = new_doc(
	"Upande Webstore — Customisation Manual",
	"Configuring the storefront and customer portal for a client project",
)

p(doc, "Every colour, word, image and feature that differs between client projects lives in Webstore Settings, not in code. Setting up a new project means filling in a form — one codebase serves any number of clients. This manual covers all five tabs.")

h(doc, "1. The one rule", 1)
p(doc, "Every field is optional, and blank means “use the shipped default”.", bold_lead="Remember this above all: ")
p(doc, "A site with nothing filled in looks exactly like the stock Ink & Gold design — the page sends no colour overrides at all. You can therefore adopt this on a live site and nothing changes until you deliberately fill something in. Clear a field again and it returns to the default.")
p(doc, "Everything else follows from that rule: you never have to fill in a whole palette to change one colour, and you can always return to a known-good state by emptying a field or applying a preset.")

h(doc, "2. Setting up a new project", 1)
p(doc, "Start from a preset and adjust — it is far faster than filling in seventy fields by hand.")
num(doc, "Open Webstore Settings → Transfer. Choose a preset close to the client (upande for the stock gold look, mona_flowers for a navy single-product-line farm) and use Theme → Apply Preset.")
num(doc, "Re-upload the images. Presets carry image links, not the files themselves; anything that did not come across is listed for you when the preset is applied. Upload the logo, favicon, hero photo and card images on the Branding tab.")
num(doc, "On the Theme tab set Accent to the client's brand colour. If that colour should also paint buttons and navigation, tick Accent Drives Primary Actions.")
num(doc, "On the Branding tab rewrite the wordmark, hero copy, footer contact details and copyright, then fill the three tables — hero stats, category cards and footer links.")
num(doc, "On the Features tab untick anything the client did not buy. Unticking hides the interface, makes the page return “not found”, and refuses the underlying request.")
num(doc, "Save, then use Theme → Export Theme and keep the downloaded file with the client's project notes. It is how you rebuild this site or clone it to staging.")
p(doc, "A brand-new installation starts on the mona_flowers preset. Upgrading an existing site never changes its appearance — existing settings are left exactly as they are.", bold_lead="Installation behaviour: ")

h(doc, "3. Theme tab — colour, type and shape", 1)
p(doc, "You set a handful of source colours and the rest of the palette is worked out from them. That is why there are thirteen colour fields rather than thirty-five: setting Ink alone produces the whole range of greys, the shadows and the dark gradients, all tinted to match.")

h(doc, "3.1 Brand colours", 2)
bullet(doc, "highlights, price chips, badges, chart series and hero italics. Hover, light, deep and focus-ring shades are derived from it.", bold_lead="Accent — ")
bullet(doc, "the deep end of the accent, used for gradients and accent-coloured text on white. Blank derives a darkened accent.", bold_lead="Accent Dark — ")
bullet(doc, "the pale tint behind accent badges and chips. Blank derives a near-white tint.", bold_lead="Accent Soft — ")
bullet(doc, "headings and body text, and the source of the whole grey range plus every shadow.", bold_lead="Ink / Neutral — ")
bullet(doc, "secondary text. Set this to control how warm or cool the greys look; left blank they are a straight blend and can look flat.", bold_lead="Muted Text — ")
bullet(doc, "the page background, and the slightly lighter surface tone above it.", bold_lead="Page Canvas — ")
bullet(doc, "sunken fills such as table row hovers and quiet panels. Should be darker than the canvas.", bold_lead="Muted Fill — ")
bullet(doc, "hairlines between cards, rows and sections.", bold_lead="Border — ")
bullet(doc, "heavier dividers and outlined-button borders.", bold_lead="Border (strong) — ")

h(doc, "3.2 Does the brand colour paint the buttons?", 2)
p(doc, "This is the most consequential choice on the tab and it has its own checkbox: Accent Drives Primary Actions.")
bullet(doc, "the accent is trim only — badges, the bold half of the wordmark, the basket count. Buttons, the active navigation pill and avatars stay near-black. This is the stock Ink & Gold look.", bold_lead="Unticked — ")
bullet(doc, "buttons, the active navigation pill, avatars and focus rings all take the accent, while headings and body text stay ink. Use this when the client's brand colour is their action colour.", bold_lead="Ticked — ")
p(doc, "If you set an Accent and the buttons stubbornly stay black, this checkbox is the reason.", bold_lead="Most common confusion: ")

h(doc, "3.3 Status, typography and shape", 2)
p(doc, "Success, Warning, Danger and Info each fill out a whole family: the flat colour, a darker or brighter partner, and a soft translucent tint for badge backgrounds. Set Warning to the darker text shade — the brighter fill is derived from it.")
p(doc, "Body, Display and Mono are dropdowns. Choose a bundled family (Poppins, Fraunces, IBM Plex Mono) and you are done — those ship with the app and need no internet connection. For anything else choose Custom, type the family name exactly as the foundry spells it, and paste a stylesheet link into Google Fonts URL. Both parts are required; a Custom choice without a family name is rejected on save.")
p(doc, "the font address must begin https://fonts.googleapis.com. Any other host is refused, because this field adds a tag to every page on the site.", bold_lead="Security note: ")
p(doc, "Base Radius, Card Radius and Panel Radius accept any CSS length (4px, 0.75rem, 0 for square corners). Defaults are 0.75rem, 20px and 24px.")

h(doc, "3.4 Custom CSS", 2)
p(doc, "Under Advanced. Applied last, so it overrides any value the fields above produce — the escape hatch for pinning one specific shade rather than fighting the derivation:")
code(doc, "--ws-radius-card: 2px;\n--ws-ink-4: #54586b;")

h(doc, "4. Branding tab — words and pictures", 1)
p(doc, "Every visible string on the storefront and portal is a field here; nothing is hardcoded.")
bullet(doc, "Site Name, Wordmark (the light half of the logo text), Wordmark (bold) (the accented half), Wordmark Subtitle, Brand Logo and Favicon.", bold_lead="Identity — ")
bullet(doc, "Hero Image, Eyebrow, Heading, Heading (emphasis) for the italic tail of the headline, Body, and three button labels covering the primary action, guests, and signed-in members.", bold_lead="Hero — ")
bullet(doc, "Tagline, Contact Email, Hours, Location, Website, Copyright Holder (the year is added automatically) and Footer Note.", bold_lead="Footer — ")
bullet(doc, "Portal Eyebrow, the small label above every customer portal page title.", bold_lead="Portal — ")
p(doc, "the guest button label is skipped when Signup is switched off; guests are shown “Member login” instead, so the button can never point at a page that is not there.", bold_lead="Interaction: ")

h(doc, "4.1 The three tables", 2)
p(doc, "Repeating content is a table, so you control how many rows there are.")
bullet(doc, "Value and Label — the figures under the hero text. Any number of rows; three fits the layout best.", bold_lead="Hero Stats — ")
bullet(doc, "Label, Subtitle, Image, Category and Custom URL. Set Category to filter the shop by it; Custom URL overrides that and can point anywhere.", bold_lead="Category Cards — ")
bullet(doc, "Column, Label and URL. Rows sharing a Column heading are grouped under it in table order, so the number of footer columns follows your rows.", bold_lead="Footer Links — ")
p(doc, "an empty table hides its section entirely. No stats, no stat row; no cards, no card grid — not an empty gap where one used to be.", bold_lead="Important: ")

h(doc, "5. Features tab — nineteen switches", 1)
p(doc, "All nineteen start switched on. Unticking one does three things at once: hides it from the interface, makes its page return “not found”, and refuses the underlying request. There is no leftover address a customer can still reach.")

h(doc, "5.1 Storefront", 2)
bullet(doc, "off gives a browse-only catalogue — no basket, no add-to-cart, no checkout. Useful when the client takes orders by phone.", bold_lead="Cart & Checkout — ")
bullet(doc, "off removes saved items and the heart buttons.", bold_lead="Wishlist — ")
bullet(doc, "off removes self-registration; guests see Member login instead of the trade-account button.", bold_lead="Signup — ")
bullet(doc, "off removes the quick search box and its keyboard shortcut. The sidebar category search stays.", bold_lead="Search Palette — ")
bullet(doc, "off sends the basket link to the full page instead of sliding out a panel.", bold_lead="Cart Drawer — ")
bullet(doc, "off opens the page straight into the catalogue.", bold_lead="Hero — ")
bullet(doc, "off keeps the hero but drops the figures underneath.", bold_lead="Hero Stats — ")
bullet(doc, "off removes the card grid above the catalogue.", bold_lead="Category Cards — ")
bullet(doc, "off removes the site footer.", bold_lead="Footer — ")

h(doc, "5.2 Customer portal", 2)
bullet(doc, "the master switch. Off closes every portal page and hides all portal links, whatever the individual switches say.", bold_lead="Portal — ")
bullet(doc, "Dashboard, Orders, Invoices, Statement, Support, Claims and Account each close their own page.", bold_lead="Per-page switches — ")
bullet(doc, "off stops customers viewing or accepting quotations online. Checkout still creates them for your team.", bold_lead="Quotations — ")
bullet(doc, "off keeps the sidebar navigation but drops the “at a glance” figures.", bold_lead="Sidebar Stats — ")
p(doc, "switching the master Portal off leaves the individual switches untouched, so turning it back on restores exactly the arrangement you had. Cart Drawer is the one exception: it switches itself off when Cart & Checkout does, because it would have nothing to show.", bold_lead="Dependencies: ")

h(doc, "6. Transfer tab — moving a theme between sites", 1)
bullet(doc, "downloads every Theme, Branding and Features value as one file. Your project backup, and how you clone a live site to staging.", bold_lead="Export Theme — ")
bullet(doc, "attach a file to Theme JSON, then press this. It asks for confirmation first.", bold_lead="Import Theme — ")
bullet(doc, "choose from Preset and press. Same as an import, using a configuration that ships with the app.", bold_lead="Apply Preset — ")
p(doc, "anything the file does not mention is reset to its default. That is deliberate: swapping between two clients' themes leaves nothing behind from the previous one. Your General settings are never touched — company, price list, warehouses and notification addresses all survive an import.", bold_lead="Importing replaces, it does not merge: ")
p(doc, "a theme file records where each image was, not the image itself, otherwise the files would be enormous. After importing you get a list of anything that does not exist on this site, and you re-upload those on the Branding tab. Images that ship with the app always resolve; ones you uploaded yourself will not cross to a different site.", bold_lead="Images are links, not copies: ")

h(doc, "7. The two shipped presets", 1)
p(doc, "Both are ordinary theme files, so they double as worked examples — apply one, look at what it filled in, and adjust from there.")
bullet(doc, "the stock Ink & Gold design (accent #d9a514, ink #0a0a0a, canvas #f4f3ef) with black buttons and the Flowers / Coffee / Fresh Produce range. Use it to return a site to the shipped appearance, or to start a multi-category grower.", bold_lead="upande — ")
bullet(doc, "navy on cool grey (accent #1e4d8c, ink #1a1a1a, canvas #f7f8fa) with navy buttons and navigation, a Roses and Eucalyptus range, and signup switched off. Use it to start a single-product-line farm with a non-black brand colour.", bold_lead="mona_flowers — ")
p(doc, "Mona's logo lives on Mona's own site, so applying that preset elsewhere reports the logo as missing. That is the importer doing its job — upload the client's own logo.", bold_lead="Expected on import: ")

h(doc, "8. Recipes", 1)
bullet(doc, "switch off Cart & Checkout. The basket, add-to-cart buttons and checkout disappear; prices and stock still show. Leave Portal on if customers should still see invoices and statements.", bold_lead="A catalogue with no online ordering — ")
bullet(doc, "switch off Hero, Category Cards and Cart & Checkout, and set the home page to /portal in Website Settings. The catalogue stays reachable but stops being the front door.", bold_lead="A customer portal with no shop — ")
bullet(doc, "export the current theme first, as your way back. Change Accent, Ink and Page Canvas, tick Accent Drives Primary Actions if the new brand colour should carry the buttons, and set Muted Text if the greys need to shift with it.", bold_lead="Rebranding a live site — ")
bullet(doc, "Export Theme on live, Import Theme on staging, then re-upload the images it lists. Staging keeps its own company and price-list settings, because an import never touches the General tab.", bold_lead="Cloning a site to staging — ")
bullet(doc, "set the seeds first and check the result. If one specific shade is still slightly off, pin just that one in Custom CSS rather than fighting the other fields.", bold_lead="Matching a brand exactly — ")

h(doc, "9. If something looks wrong", 1)
bullet(doc, "save the form, then reload the storefront — changes apply on the next page load, not live. If it still looks unchanged, force a browser refresh. Check you edited the field you meant: Accent changes trim, Ink changes text and buttons.", bold_lead="I changed a colour and the site looks the same — ")
bullet(doc, "tick Accent Drives Primary Actions. Without it the accent is decorative only and near-black keeps painting the buttons and navigation.", bold_lead="I set an Accent but the buttons are still black — ")
bullet(doc, "set Muted Text. Left blank the greys are a straight blend between ink and canvas, which comes out fairly neutral; setting it fixes the most visible grey on the site and pulls the rest into line.", bold_lead="The greys look flat, too warm, or too cold — ")
bullet(doc, "those come from tables, and an empty table hides its section. Add rows on the Branding tab, or apply a preset to get a working set you can edit.", bold_lead="My category cards or footer links vanished — ")
bullet(doc, "expected: theme files carry image links, not the files. The import lists which ones are missing here; re-upload those on the Branding tab.", bold_lead="An imported theme has broken images — ")
bullet(doc, "check two switches on the Features tab: the master Portal switch and the page's own. Either being off closes the page. This is also what a customer sees, so it is working as intended.", bold_lead="A portal page says “not found” — ")
bullet(doc, "custom fonts need both parts: the family name spelled exactly as the foundry does, and a https://fonts.googleapis.com stylesheet link. Any other address is refused on save.", bold_lead="My custom font did not load — ")
bullet(doc, "apply the upande preset for the shipped appearance, or clear the Theme fields to send no overrides at all. Either way you are back to a known state without touching code.", bold_lead="I want to start over — ")

doc.save(os.path.join(HERE, "04-customisation-manual.docx"))
print("saved 04-customisation-manual.docx")
